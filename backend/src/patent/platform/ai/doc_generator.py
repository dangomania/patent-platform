"""
Generate a formatted OA translation Word document based on OA-trans-template.docx.

Flow:
  1. Call Claude to parse the English translation into structured JSON
     (header fields, per-reason type / claims / cited docs / body text)
  2. Build a new .docx using python-docx, following the template's layout
"""

import io
import json
import re
import anthropic
import os
import os as _os

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_client = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

_TEMPLATE_PATH = _os.path.join(_os.path.dirname(__file__), "templates", "OA-trans-template.docx")

# ── Boilerplate sentences for each reason type ───────────────────────────────

_BOILERPLATE: dict[str, str] = {
    "Unity of Invention": (
        "The subject application does not satisfy the requirements as provided in "
        "Art. 37 of the Patent Law in the following respect(s)."
    ),
    "Industrial Applicability": (
        "The subject application does not satisfy the requirements as provided in "
        "Art. 29, para. 1, First sentence of the Patent Law in the following respect(s)."
    ),
    "Novelty": (
        "Claims {claims} of the application are deemed to be the same as the disclosures "
        "in the following publications."
    ),
    "Inventive Step": (
        "Claims {claims} of the application are deemed to be easily conceivable by those "
        "skilled in the art on the basis of the following publications."
    ),
    "Prior Art Effect": (
        "Claims {claims} of the application are deemed to be the same as the invention "
        "disclosed in the specification or drawings originally attached to the application "
        "of the patent application listed below."
    ),
    "Enablement Requirement": (
        "The subject application does not satisfy the requirements as provided in "
        "Art. 36, (4)-(i) of the Patent Law in the following respect(s)."
    ),
    "Written Description Requirement": (
        "The subject application does not satisfy the requirements as provided in "
        "Art. 36, (6)-(i) of the Patent Law in the following respect(s)."
    ),
    "Clarity Requirement": (
        "The subject application does not satisfy the requirements as provided in "
        "Art. 36, (6)-(ii) of the Patent Law in the following respect(s)."
    ),
    "Double Patent": (
        "The claimed invention of the captioned application is deemed to be the same as "
        "the invention of the patent/patent application listed below."
    ),
    "New Matter": (
        "The amendment filed is not made within the scope disclosed in the translations "
        "of the specification, claims or drawings originally attached to the application."
    ),
}

_OPENING = (
    "The captioned application should be rejected due to the following reasons. "
    "Should the applicant have any argument against this rejection, please submit a "
    "written opinion to the Japan Patent Office within 60 days from the mailing date "
    "of this notification."
)

_FOOTER_DIVIDER = "-" * 100
_SEARCH_RECORD_HEADER = "<Record of Search Results of Prior Art Documents>"
_AMENDMENT_NOTE_HEADER = "<Note in case of making an amendment>"

_AMENDMENT_NOTES = [
    "(1)\tWhen the applicant makes amendments to specification and claims, they should "
    "underline the descriptions modified by the amendments.",
    "(2)\tAmendments should be made within the scope of matters described in the "
    "translations of the specification, claims or drawings originally attached to the "
    "application.",
    "(3)\tThe applicant should take care not to violate Art. 17bis (4) of the Patent Law "
    "if they make amendments to claims.",
]

# ── Claude: extract structure from translation ────────────────────────────────

_EXTRACT_PROMPT = """\
You will receive an English translation of a Japanese Patent Office (JPO) Office Action.
Extract the structured information and return ONLY valid JSON — no markdown, no prose.

JSON schema:
{
  "oa_number": "1st Office Action",   // "1st", "2nd", "3rd" or "Final"
  "date": "Month DD, YYYY",           // Date field (blank if not found)
  "mailing_date": "Month DD, YYYY",   // Mailing Date (blank if not found)
  "examiner": "SURNAME, Given",       // Examiner name (blank if not found)
  "examiner_division": "",            // e.g. "Examination 3rd division, life technologies (PA4B)"
  "examiner_tel": "",                 // Tel number
  "examiner_email": "",               // Email
  "is_final": false,                  // true if Final Rejection or contains "final notification"
  "reasons": [
    {
      "number": 1,
      "type": "Novelty",              // One of: Unity of Invention, Industrial Applicability,
                                      //   Novelty, Inventive Step, Prior Art Effect,
                                      //   Enablement Requirement, Written Description Requirement,
                                      //   Clarity Requirement, Double Patent, New Matter
      "notes": [                      // One entry per NOTE block
        {
          "label": "NOTE",            // "NOTE", "NOTE 1", "NOTE 2", etc.
          "claims": "1, 2, 4-7",      // claims list after "Re.:"
          "cited_docs": [             // list of cited document strings
            "D1: VACCINES 2018..."
          ],
          "body": "Full body text of this NOTE block, preserving paragraphs with \\n\\n"
        }
      ]
    }
  ]
}

Rules:
- Extract every NOTE block separately within each REASON
- cited_docs: include each document on its own line as "DX: description"
- body: the substantive translated text (exclude NOTE header, Re.: line, cited docs lines)
- If a field is not found in the text, use empty string or empty array
- Return ONLY the JSON object, nothing else
"""


def structure_translation(translation: str) -> dict:
    """Use Claude to parse translation into structured JSON."""
    msg = _client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=4096,
        system=_EXTRACT_PROMPT,
        messages=[{"role": "user", "content": translation}],
    )
    raw = msg.content[0].text.strip()
    # Strip markdown fences if present
    raw = re.sub(r'^```[a-z]*\n?', '', raw)
    raw = re.sub(r'\n?```$', '', raw)
    return json.loads(raw)


# ── python-docx helpers ───────────────────────────────────────────────────────

def _set_font(run, size_pt: float = 11, bold: bool = False):
    run.font.name = "Aptos"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # Set East Asian font too
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Aptos')


def _add_para(doc: Document, text: str, bold: bool = False,
              style: str = "Normal", align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    run = p.add_run(text)
    _set_font(run, bold=bold)


def _add_tab_field(doc: Document, label: str, value: str) -> None:
    """Add a header field line like 'Date:\tJune 30, 2026'."""
    p = doc.add_paragraph(style="Normal")
    r = p.add_run(f"{label}\t{value}")
    _set_font(r)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_oa_docx(
    translation: str,
    app_number: str = "",
    oa_date: str = "",
) -> bytes:
    """
    Generate a formatted OA translation .docx.
    Returns raw bytes of the Word file.
    """
    # 1. Extract structure via Claude
    try:
        data = structure_translation(translation)
    except Exception:
        # Fallback: plain document with just the translation text
        data = {"reasons": [], "oa_number": "", "date": oa_date,
                "mailing_date": "", "examiner": "", "examiner_division": "",
                "examiner_tel": "", "examiner_email": "", "is_final": False}

    # 2. Build document
    doc = Document()

    # Page margins (match template ~2.5cm)
    for section in doc.sections:
        section.top_margin    = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin   = Pt(90)
        section.right_margin  = Pt(72)

    # ── Title block ──
    _add_para(doc, "TRANSLATION OF", bold=True)
    oa_type = data.get("oa_number", "")
    title = "NOTICE OF REJECTION"
    if data.get("is_final"):
        title = "FINAL REJECTION"
    _add_para(doc, title, bold=True)
    doc.add_paragraph()  # blank
    if oa_type:
        _add_para(doc, f"({oa_type})")
    doc.add_paragraph()  # blank

    # ── Header fields ──
    date_val   = data.get("date") or oa_date or ""
    mail_val   = data.get("mailing_date", "")
    exam_val   = data.get("examiner", "")
    app_val    = app_number or ""

    _add_tab_field(doc, "Date:",         date_val)
    _add_tab_field(doc, "Mailing Date:", mail_val)
    _add_tab_field(doc, "Examiner:",     exam_val)
    doc.add_paragraph()
    _add_para(doc, f"Japanese Patent Application No. {app_val}")
    doc.add_paragraph()

    # Final notification note
    if data.get("is_final"):
        _add_para(doc, "<< Final Notification of Reasons for Rejection >>")
        doc.add_paragraph()

    # ── Opening sentence ──
    _add_para(doc, _OPENING)
    doc.add_paragraph()

    # ── Reasons ──
    reasons = data.get("reasons", [])
    if not reasons:
        # No structured reasons found — just dump the translation
        for para in translation.split("\n\n"):
            if para.strip():
                _add_para(doc, para.strip())
    else:
        for reason in reasons:
            rnum  = reason.get("number", "")
            rtype = reason.get("type", "")
            notes = reason.get("notes", [])

            # Reason header
            _add_para(doc, f"REASON {rnum} ({rtype})", bold=True)

            # Boilerplate sentence
            boiler = _BOILERPLATE.get(rtype, "")
            if boiler:
                claims_str = notes[0].get("claims", "") if notes else ""
                boiler = boiler.replace("{claims}", claims_str)
                _add_para(doc, boiler, style="Normal")

            # NOTE blocks
            for note in notes:
                label     = note.get("label", "NOTE")
                claims    = note.get("claims", "")
                cited     = note.get("cited_docs", [])
                body_text = note.get("body", "")

                _add_para(doc, label)
                if claims:
                    _add_para(doc, f"Re.: Claims {claims}")
                if cited:
                    first = True
                    for cd in cited:
                        prefix = "Cited document" if len(cited) == 1 else ("Cited documents" if first else "")
                        if first:
                            _add_para(doc, f"{prefix}:\t{cd}")
                        else:
                            _add_para(doc, f"\t{cd}")
                        first = False

                # Body paragraphs
                if body_text:
                    for para in body_text.split("\n\n"):
                        para = para.strip()
                        if para:
                            _add_para(doc, para)

            doc.add_paragraph()

    # ── Footer ──
    _add_para(doc, _FOOTER_DIVIDER)
    _add_para(doc, _SEARCH_RECORD_HEADER)
    doc.add_paragraph()
    _add_para(doc, "This Record of Search Results of Prior Art Documents is not a part of NOTICE OF REJECTION")
    _add_para(doc, _FOOTER_DIVIDER)

    _add_para(doc, _AMENDMENT_NOTE_HEADER)
    for note in _AMENDMENT_NOTES:
        _add_para(doc, note)
    _add_para(doc, _FOOTER_DIVIDER)

    # Contact info
    division = data.get("examiner_division", "")
    tel      = data.get("examiner_tel", "")
    email    = data.get("examiner_email", "")
    _add_para(doc,
        "If you have any questions or if you wish to have an interview, please contact "
        "the examiner as follows."
    )
    if division: _add_para(doc, division)
    if exam_val: _add_para(doc, f"Examiner: {exam_val}")
    if tel:      _add_para(doc, f"Tel: {tel}")
    if email:    _add_para(doc, email)

    # 3. Serialise to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
